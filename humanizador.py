code_content = """import streamlit as st
import google.generativeai as genai
from docx import Document
import io

# Configuración de la página de Streamlit
st.set_page_config(page_title="Humanizador de Word", page_icon="📝", layout="centered")

st.title("📝 Humanizador y Corregidor de Documentos Word")
st.write("Sube tu archivo de Word, procésalo con IA para que suene más natural y descárgalo corregido al instante.")

# Barra lateral para la API Key de Gemini
st.sidebar.header("Configuración")
api_key = st.sidebar.text_input("Introduce tu Gemini API Key:", type="password", help="Obtenla gratis en Google AI Studio")

# Enlace de ayuda para obtener la API Key
st.sidebar.markdown("[¿Cómo obtener una API Key?](https://aistudio.google.com/)")

# Subida del archivo
uploaded_file = st.file_uploader("Elige un archivo de Word (.docx)", type=["docx"])

if uploaded_file:
    if not api_key:
        st.warning("⚠️ Por favor, introduce tu API Key en la barra lateral para continuar.")
    else:
        # Configurar la API de Google Gemini
        genai.configure(api_key=api_key)
        
        try:
            # Leer el archivo Word cargado
            doc = Document(uploaded_file)
            texto_completo = []
            for parrafo in doc.paragraphs:
                if parrafo.text.strip():
                    texto_completo.append(parrafo.text)
            
            texto_original = "\\n\\n".join(texto_completo)
            
            # Mostrar vista previa del texto original
            st.subheader("📄 Vista previa del texto original")
            st.text_area("Texto detectado:", texto_original, height=200, disabled=True)
            
            # Botón para procesar
            if st.button("✨ Humanizar y Corregir Texto"):
                with st.spinner("La IA está reescribiendo tu texto para hacerlo más humano y fluido..."):
                    
                    # Modelo a utilizar
                    modelo = genai.GenerativeModel('gemini-1.5-flash')
                    
                    # Prompt de optimización
                    prompt = f\"\"\"
                    Actúa como un editor y escritor humano experto. Toma el siguiente texto y reescríbelo 
                    para que suene mucho más natural, fluido, conversacional y empático, eliminando cualquier 
                    rastro de redacción robótica, monótona o típica de IA.
                    
                    Reglas críticas:
                    1. Varía la longitud de las oraciones (combina frases cortas e impactantes con otras más largas).
                    2. Utiliza conectores naturales y transiciones suaves.
                    3. Mantén exactamente la misma información histórica, técnica o argumental; solo cambia la FORMA en que se cuenta.
                    4. No agregues introducciones ni explicaciones tuyas, devuelve únicamente el texto corregido.
                    
                    Texto a humanizar:
                    {texto_original}
                    \"\"\"
                    
                    # Generar contenido
                    respuesta = modelo.generate_content(prompt)
                    texto_humanizado = respuesta.text
                    
                    # Mostrar el resultado en pantalla
                    st.subheader("🎉 Texto Humanizado")
                    st.text_area("Resultado:", texto_humanizado, height=200)
                    
                    # Crear el nuevo documento Word en memoria para la descarga
                    doc_salida = Document()
                    parrafos_nuevos = texto_humanizado.split('\\n\\n')
                    for p in parrafos_nuevos:
                        doc_salida.add_paragraph(p.strip())
                    
                    # Guardar en un buffer de bytes
                    buffer = io.BytesIO()
                    doc_salida.save(buffer)
                    buffer.seek(0)
                    
                    # Botón de descarga
                    st.success("¡Texto procesado con éxito! Haz clic abajo para descargarlo.")
                    st.download_button(
                        label="📥 Descargar Word Corregido",
                        data=buffer,
                        file_name="documento_humanizado.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
        except Exception as e:
            st.error(f"Ocurrió un error al procesar el archivo: {e}")
"""

with open("app_humanizador.py", "w", encoding="utf-8") as f:
    f.write(code_content)

print("Archivo creado exitosamente.")